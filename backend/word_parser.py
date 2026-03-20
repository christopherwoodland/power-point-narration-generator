"""
Parse a Word document and extract per-slide text blocks.

Convention: paragraphs styled as "Heading 1" delimit slides.
Fallback: if no Heading 1 styles are found, look for paragraphs whose text
starts with "Slide " (case-insensitive) followed by a number.
"""
import re
from docx import Document


_SLIDE_RE = re.compile(r"^slide\s+\d+", re.IGNORECASE)


def extract_slides(docx_bytes: bytes) -> list[dict]:
    """
    Returns a list of dicts:
        [{"title": "Slide 1 Welcome", "text": "Welcome to this course..."}, ...]
    The list is ordered by appearance in the document.
    Slides with no body text are included with an empty string.
    """
    import io
    doc = Document(io.BytesIO(docx_bytes))

    # Determine delimiter strategy
    has_heading1 = any(
        p.style.name == "Heading 1" for p in doc.paragraphs if p.text.strip()
    )

    slides = []
    current_title = None
    current_lines = []

    def _flush():
        if current_title is not None:
            slides.append({
                "title": current_title,
                "text": "\n".join(current_lines).strip(),
            })

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        is_delimiter = False
        if has_heading1:
            is_delimiter = (para.style.name == "Heading 1")
        else:
            is_delimiter = bool(_SLIDE_RE.match(text))

        if is_delimiter:
            _flush()
            current_title = text
            current_lines = []
        else:
            if current_title is not None:
                current_lines.append(text)

    _flush()
    return slides
